import sys
import os

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue 600
    return p

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Microsoft JhengHei"
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.name = "Microsoft JhengHei"
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_bullet_point(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r_b = p.add_run(bold_title)
    r_b.font.name = "Microsoft JhengHei"
    r_b.font.size = Pt(10)
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(15, 23, 42)
    r_t = p.add_run(text)
    r_t.font.name = "Microsoft JhengHei"
    r_t.font.size = Pt(10)
    r_t.font.color.rgb = RGBColor(51, 65, 85)
    return p

def generate_doc():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("AI Trend Listening 4.0")
    r_title.font.name = "Microsoft JhengHei"
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    r_sub = p_sub.add_run("系統完整技術架構圖與模組資料流流程圖規格書 (Mermaid Architecture Diagram Specification)")
    r_sub.font.name = "Microsoft JhengHei"
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(37, 99, 235)

    # Document Metadata Table Box
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("系統名稱 (System Name)", "AI Trend Listening 4.0 (原 智企前瞻 AI Pulse)"),
        ("文件主題 (Topic)", "4-Stage 管道架構圖、品質過濾流程圖、4-Tier 評分機制圖與雙排程架構"),
        ("文件版本 (Version)", "v4.0 (2026-07-29 最新版)"),
        ("專案負責人 (Owner)", "Pei-Ying (AI Trend Listening 團隊)"),
        ("繪圖語言 (Diagram Format)", "Mermaid Flowchart Syntax (相容 GitHub, Notion, Draw.io, Mermaid Live Editor)")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "F8FAFC")
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = "Microsoft JhengHei"
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(15, 23, 42)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = "Microsoft JhengHei"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # 1. 高階總體系統架構圖
    add_heading_1(doc, "一、 高階總體系統架構圖 (High-Level Architecture Diagram)")
    add_body_p(doc, "本圖展現 AI Trend Listening 4.0 系統從外部 RSS 新聞監聽、Stage 2 主題控制與 4-Tier 策展、Stage 3 電子報動態建置到 Stage 4 離線歸檔全流程之模組與資料庫關聯。Stage 2 擔任主題大腦，解析 data/weekly_newsletter_theme.xlsx 當前【啟用中】主題與重點領域（如：電動車, 電池管理, 智慧駕駛），動態注入 Google News 定向搜尋語法至監聽管道中：")

    mermaid_1 = """flowchart TB
    subgraph USER_LAYER["👤 使用者與維運層 (Management Layer)"]
        ADMIN["企業主管 / 電子報維運團隊"]
        CONFIG["data/weekly_newsletter_theme.xlsx\\n(每周主題與重點領域設定)"]
        ADMIN -->|發布與設定啟用中主題| CONFIG
    end

    subgraph DATA_SOURCES["📡 外部數據來源層 (Data Sources)"]
        GNEWS["Google News RSS (TW & US)\\n(8大通用語法 + 動態主題語法)"]
        RSS_FEEDS["11+ 權威科技與商業 RSS 媒體庫\\n(TechCrunch, Wired, iThome, NVIDIA 等)"]
    end

    subgraph STAGE1["📡 Stage 1: 增量新聞監聽與品質過濾器 (stage1_news_fetcher.py)"]
        CRAWLER["多源 RSS 抓取模組"]
        FILTER["3 大品質過濾演算法\\n(Negative / Use Case / Tech / Impact)"]
        DEDUP["HashSet 增量去重與限制日過濾\\n(seen_links + seen_titles, START_DATE=2026-07-27)"]
        TAXONOMY["8 大趨勢領域自動打標\\n(TREND_TAXONOMY)"]
        
        GNEWS --> CRAWLER
        RSS_FEEDS --> CRAWLER
        CRAWLER --> FILTER
        FILTER --> DEDUP
        DEDUP --> TAXONOMY
    end

    subgraph DATABASE["💾 數據持久化層 (Persistence Layer)"]
        JSON_DB[("stage1_ai_news.json\\n(全量增量數據庫)")]
        EXCEL_DB[("stage1_ai_news.xlsx\\n(多格式寫入與 Lock 備援)")]
        TAXONOMY --> JSON_DB
        TAXONOMY --> EXCEL_DB
    end

    subgraph STAGE2["🎯 Stage 2: 主題策展與 4-Tier 評分引擎 (stage2_curator.py)"]
        THEME_MGR["動態主題對接與語意擴充\\n(DOMAIN_EXPANSION)"]
        RELEVANCE_GATE["100% 主題對接熔斷機制\\n(Score = 0 直接過濾無關文章)"]
        SCORING_ENGINE["4-Tier 權重計分演算法\\n(Theme 40% + Biz 30% + Auth 20% + Appeal 10%)"]
        RATIONALE_GEN["專家評比理由與職能標籤合成"]
        
        CONFIG --> THEME_MGR
        JSON_DB --> RELEVANCE_GATE
        THEME_MGR --> RELEVANCE_GATE
        RELEVANCE_GATE -->|通過門檻| SCORING_ENGINE
        SCORING_ENGINE --> RATIONALE_GEN
    end

    subgraph STAGE2_OUTPUTS["📑 Stage 2 產出層"]
        S2_EXCEL[("stage2_curated_news.xlsx\\n(Sheet 1 權重說明 + Sheet 2 精選清單)")]
        S2_MD[("stage2_curated_report.md\\n(Markdown 策展報告)")]
        RATIONALE_GEN --> S2_EXCEL
        RATIONALE_GEN --> S2_MD
    end

    subgraph STAGE3["🎨 Stage 3: Acelia 雜誌/SaaS 電子報建置器 (stage3_newsletter_builder.py)"]
        TOP_SELECT["Top 6 案例精選與零湊數保護"]
        HTTP_DECODER["googlenewsdecoder URL 還原\\n+ HTTP 原始網頁 HTML 段落抓取"]
        DIGEST_CLASSIFIER["100% 零幻覺「四大結構標題」歸類演算法\\n(📌背景 / ⚙️技術 / 📊數據 / 🚀展望)"]
        TAKEAWAYS["當期動態 Hero Takeaways 提煉"]
        DOM_SYNC["Acelia SaaS UI 與前端數據嵌入\\n(Justified Text + EMBEDDED_CASES)"]
        
        S2_EXCEL --> TOP_SELECT
        TOP_SELECT --> HTTP_DECODER
        HTTP_DECODER --> DIGEST_CLASSIFIER
        DIGEST_CLASSIFIER --> TAKEAWAYS
        TAKEAWAYS --> DOM_SYNC
    end

    subgraph WEB_OUTPUTS["🌐 前端發行層 (Web Publishing)"]
        HTML_PAGE["newsletter.html\\n(Acelia 雜誌質感電子報網頁)"]
        JS_DATA["newsletter.js\\n(前端動態渲染與 Modal 邏輯)"]
        CSS_STYLE["newsletter.css\\n(雙邊對齊與 HSL 顏色系統)"]
        DOM_SYNC --> HTML_PAGE
        DOM_SYNC --> JS_DATA
    end

    subgraph STAGE4["📦 Stage 4: 期數歸檔與離線單檔生成引擎 (stage3_archive_engine.py)"]
        SNAPSHOT["期數資料夾全資產快照\\n(archives/Issue_XX_Date_Theme/)"]
        INLINER["HTML/CSS/JS 單檔內嵌引擎\\n(standalone_newsletter.html)"]
        HTML_PAGE --> SNAPSHOT
        JS_DATA --> SNAPSHOT
        CSS_STYLE --> SNAPSHOT
        S2_EXCEL --> SNAPSHOT
        SNAPSHOT --> INLINER
    end

    subgraph SCHEDULER["⏰ 全自動雙排程常駐服務 (master_scheduler_service.py)"]
        TASK1["任務 1: 每日 17:00 觸發 Stage 1"]
        TASK2["任務 2: 每週五 09:00 AM 觸發 Master Pipeline (Stage 2 ➔ 3 ➔ 4)"]
        TASK1 -->|每日 17:00| STAGE1
        TASK2 -->|每週五 09:00| STAGE2
    end

    subgraph READERS["👥 讀者與訂閱層 (Subscribers)"]
        WEB_READER["Web 瀏覽器讀者\\n(線上閱讀 + Modal 筆記彈窗)"]
        OFFLINE_READER["離線單檔閱讀 / Email 附件\\n(standalone_newsletter.html)"]
        HTML_PAGE --> WEB_READER
        INLINER --> OFFLINE_READER
    end"""

    add_code_block(doc, mermaid_1)

    # 2. Stage 1 流程圖
    add_heading_1(doc, "二、 Stage 1 增量新聞抓取與品質過濾流程圖 (Stage 1 Flowchart)")
    add_body_p(doc, "Stage 1 結合 Google News 8 大定向 Query、當期動態主題 Query 以及 11+ 權威 RSS 媒體庫，並透過 HashSet 與 3 大品質過濾器過濾炒作與廣告：")

    mermaid_2 = """flowchart TD
    Start([開始 Stage 1 增量新聞抓取]) --> ReadTheme[讀取 Excel 當前【啟用中】主題與重點領域]
    ReadTheme --> GenQueries[產生搜尋語法:\\n1. 8大通用 Query\\n2. 動態主題屬性 Query\\n3. 11+ 固定 RSS 網址]
    GenQueries --> LoadOldDB[載入既存 stage1_ai_news.json\\n建立 seen_links 與 seen_titles HashSet]
    LoadOldDB --> FetchRSS[爬取 Google News TW/US 與 權威 RSS]
    
    FetchRSS --> LoopArticle{逐篇檢驗文章}
    LoopArticle -->|已被 HashSet 紀錄 OR 早於 START_DATE| Skip[過濾拋棄]
    LoopArticle -->|新文章| CheckNeg{是否包含 EXCLUDE_TERMS?\\n融資炒作/活動報名廣告}
    
    CheckNeg -->|是| Skip
    CheckNeg -->|否| CheckUseCase{是否包含 USE_CASE_TERMS?\\n具體應用情境}
    
    CheckUseCase -->|否| Skip
    CheckUseCase -->|是| CheckTech{是否包含 TECH_TERMS?\\n技術與軟硬體細節}
    
    CheckTech -->|否| Skip
    CheckTech -->|是| CheckImpact{是否包含 IMPACT_TERMS?\\n量化效益與處置方案}
    
    CheckImpact -->|否| Skip
    CheckImpact -->|是| PassQuality[通過 3 大品質過濾器]
    
    PassQuality --> Taxonomy[標註 8 大趨勢標籤 TREND_TAXONOMY]
    Taxonomy --> AddToDB[加入全量新聞列表 & HashSet]
    
    Skip --> NextArticle{還有下一篇文章?}
    AddToDB --> NextArticle
    NextArticle -->|是| LoopArticle
    NextArticle -->|否| SaveData[同步寫入 JSON / CSV / Excel]
    SaveData --> End1([Stage 1 增量抓取完成])"""

    add_code_block(doc, mermaid_2)

    # 3. Stage 2 流程圖
    add_heading_1(doc, "三、 Stage 2 4-Tier 評分與 100% 熔斷門檻流程圖 (Stage 2 Flowchart)")
    add_body_p(doc, "Stage 2 實施 Strict Theme Relevance Gate 熔斷機制。非當期主題新聞 Score = 0 直接剔除，合格文章則進入 4-Tier 權重評分模組：")

    mermaid_3 = """flowchart TD
    Start2([開始 Stage 2 主題策展]) --> LoadTheme[載入 Excel 當前【啟用中】主題與領域]
    LoadTheme --> BuildMatcher[建構動態關鍵字庫與 DOMAIN_EXPANSION 同義詞庫]
    BuildMatcher --> LoadAll[載入 Stage 1 全量增量新聞庫]
    
    LoadAll --> EvalLoop{逐篇進行評分驗證}
    EvalLoop --> MatchDomains[計算與當期主題/領域之命中次數]
    
    MatchDomains --> GateCheck{命中次數 == 0 ?}
    GateCheck -->|是: 完全非當期主題| ScoreZero[Score = 0 分\\n100% Strict Relevance Gate 熔斷排除]
    
    GateCheck -->|否: 契合當期主題| CalcScore[計算 4-Tier 權重得分]
    
    subgraph FOUR_TIERS["4-Tier 權重計分機制 (最高 100 分)"]
        T1["Tier 1: 主題契合度 (40% 權重, 上限 50分)\\nmin(matched_domains*15 + matched_kws*4, 50)"]
        T2["Tier 2: 實務價值與 ROI (30% 權重, 上限 25分)\\nmin(count(biz_kws)*3, 25)"]
        T3["Tier 3: 媒體權威與時效 (20% 權重, 上限 15分)\\n權威媒體 15分 / 一般媒體 10分"]
        T4["Tier 4: 可讀性與切入點 (10% 權重, 上限 10分)\\n摘要長度 > 40字給 10分"]
    end
    
    CalcScore --> FOUR_TIERS
    FOUR_TIERS --> SumScore[加總得出 Total Score]
    SumScore --> CheckCutoff{Total Score >= 35 分?}
    CheckCutoff -->|否| Reject[排除未達門檻文章]
    CheckCutoff -->|是| Accept[選入候選名單 & 合成專家評比理由 Rationale]
    
    ScoreZero --> NextEval{還有下一篇文章?}
    Reject --> NextEval
    Accept --> NextEval
    
    NextEval -->|是| EvalLoop
    NextEval -->|否| RankTop[按分數高低排序，精選 Top 20 文章]
    RankTop --> ExportStage2[匯出 stage2_curated_news.xlsx 與 report.md]
    ExportStage2 --> End2([Stage 2 策展完成])"""

    add_code_block(doc, mermaid_3)

    # 4. Stage 3 流程圖
    add_heading_1(doc, "四、 Stage 3 HTTP 全文解碼與「四大結構標題」歸類圖 (Stage 3 Flowchart)")
    add_body_p(doc, "Stage 3 使用 googlenewsdecoder 解析原始媒體網址，並經 HTTP 發起全文抓取，100% 忠於原文歸類至四大結構標題，杜絕生成幻覺：")

    mermaid_4 = """flowchart LR
    subgraph INPUT["輸入來源"]
        S2_NEWS["Stage 2 精選 Top 6 文章"]
    end

    subgraph HTTP_ENGINE["🌐 HTTP 網頁全文解碼與抓取模組"]
        GDECODER["googlenewsdecoder\\n(new_decoderv1 解碼真實網址)"]
        REQ["urllib.request\\n(發起 HTTP 請求抓取原始 HTML)"]
        P_EXTRACT["<p> 標籤提取與廣告/Cookie 雜訊過濾"]
        S2_NEWS --> GDECODER --> REQ --> P_EXTRACT
    end

    subgraph CLASSIFIER["🧩 100% 零幻覺「四大結構標題」歸類演算法"]
        SEC1["📌 一、 核心背景與報導概要\\n(事件起因與總體摘要段落)"]
        SEC2["⚙️ 二、 關鍵技術與產品細節\\n(模型、晶片、軟硬體架構段落)"]
        SEC3["📊 三、 營運數據與市場動態\\n(財務數字、良率、營收、毛利段落)"]
        SEC4["🚀 四、 戰略佈局與產業展望\\n(長遠佈局、合作廠房、供應鏈段落)"]
        P_EXTRACT --> SEC1 & SEC2 & SEC3 & SEC4
    end

    subgraph DOM_SYNC["🎨 DOM 更新與 Acelia 前端嵌入"]
        TAKEAWAY_GEN[" Hero 三大 Takeaways 動態提煉"]
        TEXT_ALIGN["Justified Text (text-align: justify;)"]
        JS_EMBED["newsletter.js (EMBEDDED_CASES)"]
        HTML_UPDATE["newsletter.html (DOM 節點自動替換)"]
        
        SEC1 & SEC2 & SEC3 & SEC4 --> TAKEAWAY_GEN
        TAKEAWAY_GEN --> TEXT_ALIGN --> JS_EMBED --> HTML_UPDATE
    end"""

    add_code_block(doc, mermaid_4)

    # 5. 排程服務流程圖
    add_heading_1(doc, "五、 全自動雙排程守護服務流程圖 (Scheduler Flowchart)")
    add_body_p(doc, "master_scheduler_service.py 常駐執行雙排程追蹤，每日 17:00 執行新聞抓取，每週五 09:00 AM 執行主發行管道：")

    mermaid_5 = """flowchart TD
    Daemon[master_scheduler_service.py\\n每 15 秒循環監控系統時間] --> CheckTime{判斷觸發時間點}
    
    CheckTime -->|每日 17:00 & 當日未觸發| ExecTask1[執行 任務 1: stage1_news_fetcher.py]
    ExecTask1 --> FetchDaily[爬取與累積最新 AI 新聞到 stage1_ai_news.json/xlsx]
    FetchDaily --> SetStage1Done[標記 last_stage1_date = 今天]
    
    CheckTime -->|每週五 09:00 AM & 當日未觸發| ExecTask2[執行 任務 2: master_run_pipeline.py]
    ExecTask2 --> PipeS2[Stage 2: Active Theme 4-Tier 策展評分]
    PipeS2 --> PipeS3[Stage 3: Acelia 電子報建置 & HTTP 段落解碼]
    PipeS3 --> PipeS4[Stage 4: Issue 快照歸檔 & 離線單檔打包]
    PipeS4 --> SetPipelineDone[標記 last_pipeline_date = 今天]
    
    SetStage1Done --> WaitNext[等待下一次時間檢查 (15s)]
    SetPipelineDone --> WaitNext
    CheckTime -->|非觸發時間點| WaitNext
    WaitNext --> Daemon"""

    add_code_block(doc, mermaid_5)

    output_filename = "AI_Trend_Listening_系統架構圖與流程圖規格書.docx"
    try:
        doc.save(output_filename)
        print(f"✅ [SUCCESS] Saved Architecture Diagrams Word Document at: {os.path.abspath(output_filename)}")
    except PermissionError:
        fallback_name = "AI_Trend_Listening_系統架構圖與流程圖規格書_latest.docx"
        doc.save(fallback_name)
        print(f"⚠️ [Notice] Saved fallback Word Document at: {os.path.abspath(fallback_name)}")
    except Exception as e:
        print(f"⚠️ [Error] Could not save docx: {e}")

if __name__ == "__main__":
    generate_doc()
