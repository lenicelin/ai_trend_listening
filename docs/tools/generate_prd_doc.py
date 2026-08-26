import sys
import os

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue 600
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110) # Teal 700
    return p

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Microsoft JhengHei"
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.name = "Microsoft JhengHei"
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_bullet_point(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r_b = p.add_run(bold_title)
    r_b.font.name = "Microsoft JhengHei"
    r_b.font.size = Pt(10)
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(15, 23, 42)
    r_t = p.add_run(text)
    r_t.font.name = "Microsoft JhengHei"
    r_t.font.size = Pt(10)
    r_t.font.color.rgb = RGBColor(51, 65, 85)
    return p

def generate_prd():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("AI Trend Listening 4.0 (原 智企前瞻 AI Pulse)")
    r_title.font.name = "Microsoft JhengHei"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    r_sub = p_sub.add_run("商業 AI 趨勢即時監聽與自動化電子報發行系統 — 產品需求與技術深度規格書 (PRD v4.0 完整版)")
    r_sub.font.name = "Microsoft JhengHei"
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(37, 99, 235)

    # Metadata Table
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("文件版本 (Version)", "v4.0 (2026-07-29 系統最新技術架構與完整篩選機制版)"),
        ("專案負責人 (Owner)", "Pei-Ying (AI Trend Listening 團隊)"),
        ("系統品牌 (Brand Name)", "AI Trend Listening (原 智企前瞻 AI Pulse)"),
        ("系統架構 (Architecture)", "4-Stage Multistage Pipeline & Automated Scheduler Service (Stage 1 每日 17:00 / Master Pipeline 每週五 09:00 AM)"),
        ("數據與主題策展原則", "主動主題嚴格對接 (Strict Theme 100% Relevance Gate) + 增量新聞積累 (HashSet 只增不刪，起始日 2026-07-27)"),
        ("閱讀筆記規範 (Digest Standard)", "100% 忠於原文網頁全文 (零幻覺) + 四大主題結構標題 + 移除無謂模板贅字"),
        ("電子報視覺風格 (UX Style)", "Acelia 雜誌/SaaS 質感風格 (雙邊對齊 Justified Text + HSL 膠囊標籤 + Modal Drawer)")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "F8FAFC")
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = "Microsoft JhengHei"
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(15, 23, 42)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = "Microsoft JhengHei"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. 專案背景與願景
    add_heading_1(doc, "一、 專案背景與產品願景 (Background & Vision)")
    add_heading_2(doc, "1.1 產業背景與市場痛點")
    add_body_p(doc, "隨著生成式 AI (Generative AI) 與 Agentic AI (代理式 AI) 的爆發性成長，企業高階決策者、營運主管與廣大職場員工面臨嚴重的「資訊過載」與「真假資訊混雜」問題。市面上多數 AI 科技新聞充斥著純融資炒作、股價波動、公關宣傳或活動報名廣告，缺乏能真正指導企業轉型、具備量化效益 (Operational ROI) 與技術細節的落地實務內容。")
    add_heading_2(doc, "1.2 產品定位與核心願景")
    add_body_p(doc, "《AI Trend Listening》（原 智企前瞻 AI Pulse）旨在打造一套全自動、高可靠性的商業 AI 趨勢監聽與電子報發行系統。透過『增量新聞監聽 ➔ 4-Tier 智慧主題策展 ➔ Acelia 雜誌/SaaS 質感電子報建置 ➔ 自動化期數歸檔』全流程 Pipeline，每日 17:00 自動為企業同仁與高階主管提供最具商業價值與職場落地指引的趨勢情報。")

    # 2. 目標使用者與關鍵痛點
    add_heading_1(doc, "二、 目標使用者與關鍵痛點 (Target Persona & Key Pain Points Solved)")
    add_heading_2(doc, "2.1 目標使用者畫像 (Target Persona)")
    add_bullet_point(doc, "企業職場同仁與工程師 (Enterprise Employees & Engineers)：", "希望了解最新 AI 工具在日常工作流、智慧製造、供應鏈運籌與專案協作中的實務落地應用。注重具體技術細節、軟硬體規格與可執行的操作指引。")
    add_bullet_point(doc, "高階決策者與部門主管 (Executives & Operations Managers)：", "需要權威媒體來源、量化財務與營運數據（如 %、成、倍、億元、ROI），作為年度數位轉型與算力佈局決策依據。注重專家評比切入點、戰略佈局與風險控管指引。")

    add_heading_2(doc, "2.2 核心痛點與系統性解法矩陣")
    table_pains = doc.add_table(rows=5, cols=3)
    table_pains.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["核心痛點", "傳統新聞/報表缺點", "AI Trend Listening v4.0 系統性解法"]
    hdr_row = table_pains.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "1E293B")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pains_data = [
        ("痛點 1 — 資訊噪音與主題脫節", "報導大量無關醫療資安、股票炒作或論壇報名廣告。", "Stage 1 3大品質過濾器 + Stage 2 100% 主題對接熔斷機制 (Score = 0 直接排除)。"),
        ("痛點 2 — 摘要空泛與生成幻覺", "LLM 自行生成摘要容易產生虛構數據或機械式廢話。", "Stage 3 HTTP 網頁全文解碼與真實段落抓取，100% 忠於原文內文，杜絕幻覺。"),
        ("痛點 3 — 閱讀體驗缺乏結構", "段落混亂、字句未對齊、缺乏層次感。", "四大結構化主題標題 + Acelia SaaS 前端雙邊對齊 (Justified Text) 專業視覺。"),
        ("痛點 4 — 期數歷史資產散失", "發行後舊內容無法追溯、離線無法閱讀。", "Stage 4 自動化獨立期數資料夾歸檔 + 單檔內嵌 HTML (standalone_newsletter.html)。")
    ]
    for row_idx, row_data in enumerate(pains_data, start=1):
        row = table_pains.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            bg_color = "F1F5F9" if col_idx == 0 else "F8FAFC"
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Microsoft JhengHei"
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 23, 42)
            else:
                r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. 系統總體技術架構
    add_heading_1(doc, "三、 系統總體技術架構 (System Architecture Overview)")
    add_heading_2(doc, "3.1 4-Stage 多階管道 Pipeline 數據流")
    add_body_p(doc, "系統採用 Modular Multi-Stage Pipeline 架構，各 Stage 職責分離並透過標準 JSON/Excel 檔案進行資料解耦與解碼傳遞： Stage 1 新聞監聽庫 ➔ Stage 2 4-Tier 策展 ➔ Stage 3 電子報建置 ➔ Stage 4 自動歸檔。")
    add_heading_2(doc, "3.2 系統雙排程機制 (master_scheduler_service.py)")
    add_bullet_point(doc, "任務 1（每日增量新聞監聽）：", "每日 17:00 自動觸發 stage1_news_fetcher.py 進行新聞抓取與增量累積。")
    add_bullet_point(doc, "任務 2（每週電子報生成與歸檔）：", "每週五上午 09:00 AM 自動調用 master_run_pipeline.py 串接 Stage 2 ➔ Stage 3 ➔ Stage 4 完成電子報發行與快照歸檔。")

    # 4. Stage 1
    add_heading_1(doc, "四、 Stage 1 深度技術規格：增量新聞監聽與品質過濾器 (stage1_news_fetcher.py)")
    add_heading_2(doc, "4.1 增量數據保護與 HashSet 去重機制")
    add_body_p(doc, "系統貫徹『只增不刪 (Zero-Data-Loss Incremental Persistence)』原則，以 START_DATE = '2026-07-27' 為基準。啟動時預先載入 data/stage1_ai_news.json，維護 seen_links 與 seen_titles 兩個 HashSet，確保重複爬取的文章不會被再次寫入，並永久保留歷史已累積之新聞。")
    
    add_heading_2(doc, "4.2 動態主題搜尋語法對接 (get_active_theme_queries)")
    add_body_p(doc, "Stage 1 接受 Stage 2 主題控制器之導引，自動載入由 Stage 2 解析 data/weekly_newsletter_theme.xlsx 當前【啟用中】主題與重點對應領域關鍵字（如：電動車, 電池管理, 智慧駕駛）所生成之 Google News 定向 RSS 搜尋指令（如 AI {active_theme} after:2026-07-26），實現新聞監聽與當期主題的高效連動。")

    add_heading_2(doc, "4.3 3 大品質過濾演算法 (is_high_quality_article)")
    add_bullet_point(doc, "1. 負向排除條件 (Negative Filter - EXCLUDE_TERMS)：", "排除純粹資金炒作 (單純融資, 估值飆升, series a) 或公關活動報名 (論壇報名, 活動報名, 免費報名, summit, event go)。")
    add_bullet_point(doc, "2. 應用情境條件 (Specific Use Case - USE_CASE_TERMS)：", "必須包含具體業務/工業情境 (客服, 供應鏈, 維護, 自動化, 倉儲, 機房, 良率, 電動車, 自駕)。")
    add_bullet_point(doc, "3. 技術架構條件 (Technical Details - TECH_TERMS)：", "必須包含 AI 或軟硬體技術名稱 (llama, claude, gemini, gpt, deepseek, rag, agent, fine-tuning, copilot, transformer)。")
    add_bullet_point(doc, "4. 量化效益條件 (Quantifiable Impact - IMPACT_TERMS)：", "必須包含量化數據或解決方案詞彙 (%, 成, 倍, 縮短, 提升, 降低, ROI, 處置, 解決方案, 步驟, 指引)。")

    # 5. Stage 2
    add_heading_1(doc, "五、 Stage 2 深度技術規格：主題策展與 4-Tier 評分引擎 (stage2_curator.py)")
    add_heading_2(doc, "5.1 動態主題管理者與動態搜尋對接引擎 (get_active_theme_queries & DOMAIN_EXPANSION)")
    add_body_p(doc, "Stage 2 擔任系統之「主題大腦 (Theme Controller)」，負責自動解析 data/weekly_newsletter_theme.xlsx 當前【啟用中】主題與重點對應領域（如：電動車, 電池管理, 智慧駕駛），動態生成 Google News RSS 定向搜尋語法注入監聽管道，並建構 DOMAIN_EXPANSION 同義詞庫。")

    add_heading_2(doc, "5.2 嚴格主題對接門檻過濾 (100% Relevance Gate)")
    add_body_p(doc, "evaluate_article 函數設有 Strict Theme Relevance Filter 熔斷機制。當文章標題與內文對當期主題標題及重點領域關鍵字之命中次數為 0 時，直接回傳 Score = 0，100% 排除非當期主題新聞（如無關醫療資安或消費電子新聞）。")

    add_heading_2(doc, "5.3 4-Tier 權重計分演算法 (最高 100 分)")
    table_score = doc.add_table(rows=5, cols=4)
    table_score.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_headers = ["計分層級 (Tier)", "權重配比", "滿分上限", "計算公式與評估標準說明"]
    for idx, text in enumerate(s_headers):
        cell = table_score.rows[0].cells[idx]
        set_cell_background(cell, "0F766E")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    score_data = [
        ("Tier 1: 主題契合度", "40%", "50 分", "min(matched_domains * 15 + matched_kws * 4, 50)。評估與當期主題與重點領域之匹配密度。"),
        ("Tier 2: 實務價值與 ROI", "30%", "25 分", "min(count(biz_kws) * 3, 25)。檢測包含 工廠, 生產力, 良率, 成本, 縮短, %, 億 等實務數據詞。"),
        ("Tier 3: 媒體權威與時效", "20%", "15 分", "命中權威媒體 (iThome, TechCrunch, Wired, UDN, 天下雜誌, NVIDIA 等) 給 15 分，其餘給 10 分。"),
        ("Tier 4: 可讀性與切入點", "10%", "10 分", "新聞摘要長度 > 40 字且語義完整給 10 分，否則給 5 分。")
    ]
    for row_idx, row_data in enumerate(score_data, start=1):
        row = table_score.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            bg_color = "F1F5F9" if col_idx == 0 else "F8FAFC"
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Microsoft JhengHei"
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 23, 42)
            elif col_idx == 1:
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 118, 110)
            else:
                r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 6. Stage 3
    add_heading_1(doc, "六、 Stage 3 深度技術規格：Acelia 質感電子報建置器 (stage3_newsletter_builder.py)")
    add_heading_2(doc, "6.1 動態 Top-N 案例載入與零湊數機制")
    add_body_p(doc, "從 Stage 2 產出中自動載入精選文章，最多取前 6 高分案例 (Top 6)。若符相當期主題的新聞不足 6 篇（如 2 篇或 4 篇），系統動態依據實際篇數呈現，標題同步改為『精選 {case_count} 大深度實務案例』，絕不硬湊假案例。")

    add_heading_2(doc, "6.2 HTTP 網頁全文解碼與「四大結構標題」零幻覺歸類")
    add_bullet_point(doc, "googlenewsdecoder URL 還原：", "利用 new_decoderv1 將 Google News 轉址還原為媒體原始 URL。")
    add_bullet_point(doc, "HTTP 全文抓取與 100% 零幻覺歸類：", "發起 HTTP 請求抓取原始 HTML <p> 段落，根據關鍵字自動歸類為以下四大結構標題：")
    add_code_block(doc,
        "📌 一、 核心背景與報導概要 （事件起因與總體摘要）\n"
        "⚙️ 二、 關鍵技術與產品細節 （包含 AI 模型、晶片、軟硬體架構）\n"
        "📊 三、 營運數據與市場動態 （包含 財務數字、交車量、良率、營收、毛利）\n"
        "🚀 四、 戰略佈局與產業展望 （包含 長遠佈局、合作廠房、供應鏈規劃）"
    )

    add_heading_2(doc, "6.3 Acelia SaaS 前端視覺與雙邊對齊規格")
    add_body_p(doc, "前端 CSS (newsletter.css) 對導讀簡介 (.lead-in-card p)、重點速覽 (.highlight-item p)、案例摘要 (.case-summary) 及 Modal Drawer 彈窗 (#modal-digest-box) 統一套用 text-align: justify;，打造專業雜誌質感。")

    # 7. Stage 4
    add_heading_1(doc, "七、 Stage 4 深度技術規格：期數歸檔與離線單檔生成引擎 (stage3_archive_engine.py)")
    add_body_p(doc, "每次執行自動於 archives/ 建立獨立期數資料夾 (archives/Issue_{期數}_{日期}_{主題}/)，完整複製 HTML, CSS, JS, Excel 報告與 JSON 案例，並自動生成 CSS/JS 內嵌之離線單檔 standalone_newsletter.html。")

    # 8. 操作指令手冊
    add_heading_1(doc, "八、 系統操作指令手冊 (Execution Guide)")
    add_code_block(doc,
        "# 1. 執行主管道 (Stage 2 ➔ Stage 3 ➔ Stage 4 Archive Engine)\n"
        ".venv\\Scripts\\python.exe master_run_pipeline.py\n\n"
        "# 2. 手動執行 Stage 1 增量新聞抓取\n"
        ".venv\\Scripts\\python.exe stage1_news_fetcher.py\n\n"
        "# 3. 重新產生 PRD 規格書 Document (Word v4.0)\n"
        ".venv\\Scripts\\python.exe generate_prd_doc.py\n\n"
        "# 4. 啟動每日 17:00 全自動常駐排程服務\n"
        ".venv\\Scripts\\python.exe master_scheduler_service.py"
    )

    # 9. 版本變更對照
    add_heading_1(doc, "九、 PRD v4.0 版本變更對照表 (Version Change Log)")
    table_ver = doc.add_table(rows=7, cols=3)
    table_ver.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_headers = ["變更項目", "PRD v3.5 (舊版)", "PRD v4.0 (最新版)"]
    for idx, text in enumerate(v_headers):
        cell = table_ver.rows[0].cells[idx]
        set_cell_background(cell, "2563EB")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    v_data = [
        ("品牌名稱", "智企前瞻 AI Pulse", "AI Trend Listening 全面正名"),
        ("Stage 1 去重機制", "基礎 URL 去重", "HashSet (Link + Title) 增量數據保護 (START_DATE=2026-07-27)"),
        ("Stage 2 篩選邏輯", "基礎關鍵字計分", "100% 主題對接熔斷機制 (Score = 0 直接排除) + 4-Tier 權重計分"),
        ("Stage 3 摘要來源", "靜態模板卡片", "googlenewsdecoder HTTP 全文解碼 + 四大結構化段落歸類"),
        ("Stage 3 UI Takeaways", "跨期固定靜態文字", "當期精選 Top 6 案例動態提煉生成三大 Takeaways"),
        ("Stage 4 歸檔機制", "手動備份", "全自動期數資料夾快照 + standalone_newsletter.html 單檔 Inliner")
    ]
    for row_idx, row_data in enumerate(v_data, start=1):
        row = table_ver.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            bg_color = "F1F5F9" if col_idx == 0 else "F8FAFC"
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Microsoft JhengHei"
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 23, 42)
            elif col_idx == 2:
                r.font.bold = True
                r.font.color.rgb = RGBColor(37, 99, 235)
            else:
                r.font.color.rgb = RGBColor(51, 65, 85)

    output_filename = "智企前瞻_AI_Pulse_產品需求規格書_PRD_v4.docx"
    try:
        doc.save(output_filename)
        print(f"✅ [SUCCESS] Saved PRD Word Document at: {os.path.abspath(output_filename)}")
    except PermissionError:
        fallback_name = "智企前瞻_AI_Pulse_產品需求規格書_PRD_v4_latest.docx"
        doc.save(fallback_name)
        print(f"⚠️ [Notice] {output_filename} is locked by Microsoft Word. Saved to fallback: {os.path.abspath(fallback_name)}")
    except Exception as e:
        print(f"⚠️ [Notice] Could not save docx: {e}")

if __name__ == "__main__":
    generate_prd()
